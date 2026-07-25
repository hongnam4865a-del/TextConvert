#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DOCX 与 HTML 互转引擎（基于 python-docx）"""

import html as html_module
from pathlib import Path
from typing import Set

from bs4 import BeautifulSoup
from docx import Document

from engines.base import BaseEngine
from utils.file_utils import safe_read_text, safe_write_text


class DocxEngine(BaseEngine):
    """DOCX 转换引擎"""

    name = "docx"

    @property
    def supported_sources(self) -> Set[str]:
        return {"docx", "html", "htm"}

    @property
    def supported_targets(self) -> Set[str]:
        return {"docx", "html", "htm"}

    def to_html(self, src_path: Path, dest_path: Path) -> Path:
        """DOCX -> HTML"""
        doc = Document(str(src_path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            escaped = html_module.escape(text)
            if para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.replace("Heading ", ""))
                except Exception:
                    level = 2
                parts.append(f"<h{level}>{escaped}</h{level}>")
            else:
                parts.append(f"<p>{escaped}</p>")

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [f"<td>{html_module.escape(cell.text)}</td>" for cell in row.cells]
                rows.append("<tr>" + "".join(cells) + "</tr>")
            parts.append("<table>" + "".join(rows) + "</table>")

        body = "\n".join(parts)
        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{html_module.escape(src_path.stem)}</title>
    <style>body {{ font-family: Georgia, serif; line-height: 1.7; max-width: 800px; margin: 40px auto; padding: 0 20px; }}</style>
</head>
<body>
{body}
</body>
</html>
"""
        safe_write_text(dest_path, html_doc)
        return dest_path

    def from_html(self, html_path: Path, dest_path: Path) -> Path:
        """HTML -> DOCX"""
        html = safe_read_text(html_path)
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style"]):
            tag.decompose()

        doc = Document()
        for elem in soup.find_all(["h1", "h2", "h3", "p", "li"]):
            text = elem.get_text(strip=True)
            if not text:
                continue
            if elem.name == "h1":
                doc.add_heading(text, level=1)
            elif elem.name == "h2":
                doc.add_heading(text, level=2)
            elif elem.name == "h3":
                doc.add_heading(text, level=3)
            elif elem.name == "li":
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text)

        doc.save(str(dest_path))
        return dest_path
